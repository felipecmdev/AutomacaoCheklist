import re
import shutil
from docx import Document

class CriarDoc:
    def __init__(self):
        self.modelo_path = 'C:\\AutomacaoCheklist\\ModCL\\ModeloChecklist.docx'  

    def salvarDoc(self, infosParaSalvar, novo_nome_arquivo):
        # Copiar o arquivo modelo para o novo local
        shutil.copy(self.modelo_path, novo_nome_arquivo)
        
        # Abrir o novo arquivo copiado
        self.document = Document(novo_nome_arquivo)

        # Iterar sobre os itens do dicionário
        for chave, valor in infosParaSalvar.items():
            textoOG = re.compile(re.escape(chave))

            # Substituir o texto nos parágrafos
            for paragrafos in self.document.paragraphs:
                if textoOG.search(paragrafos.text):
                    inline = paragrafos.runs

                    for i in range(len(inline)):
                        if textoOG.search(inline[i].text):
                            text = textoOG.sub(valor, inline[i].text)
                            inline[i].text = text

            # Substituir o texto nas tabelas
            for table in self.document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if textoOG.search(paragraph.text):
                                paragraph.text = textoOG.sub(valor, paragraph.text)

        self.document.save(novo_nome_arquivo)